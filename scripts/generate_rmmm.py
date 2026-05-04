"""
Generate RMMM Plan document for the Digital Twin Health Activity Tracker project.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PROJECT_NAME = "Digital Twin Health Activity Tracker"
DATE = "29/04/2026"

# (ID, Risk, Category, Probability%, Impact, RMMM)
RISKS = [
    (1, "Schedule slippage causing missed semester / sprint deadlines",
     "BU", 70, 2,
     "Maintain Gantt chart with buffer time, weekly stand-ups, track velocity, "
     "re-prioritize backlog, communicate slippage to mentor early."),
    (2, "Inaccurate digital twin health-status computation (activity aggregation, scoring logic)",
     "TE", 65, 2,
     "Define formulas with mentor sign-off, write unit tests for aggregation, "
     "validate against sample dataset, peer-review algorithm before integration."),
    (3, "Limited team experience with full-stack Node.js + SQLite + Playwright testing",
     "ST", 60, 2,
     "Conduct knowledge-transfer sessions, pair-programming, follow official docs, "
     "use linters and code reviews, allocate learning sprints."),
    (4, "Breach of sensitive user health data (privacy / weak authentication on /login)",
     "CU", 55, 1,
     "Hash passwords (bcrypt), enforce HTTPS in deployment, validate inputs with "
     "express-validator, restrict DB access, add rate limiting, audit auth flow."),
    (5, "SQLite (better-sqlite3) scalability and data-corruption risk as activity records grow",
     "PS", 50, 2,
     "Add DB indexes on userId/date, run periodic VACUUM, schedule automated backups "
     "of health.db, plan migration path to PostgreSQL if record count > 100k."),
    (6, "Frequent change in requirements / scope creep from stakeholder feedback",
     "PD", 45, 3,
     "Adopt change-control process, freeze scope per sprint, log all change "
     "requests, assess impact before approval, get sign-off on SRS."),
    (7, "Native module (better-sqlite3) build failures across dev machines / Node versions",
     "DE", 35, 3,
     "Pin Node version via .nvmrc, document setup in README, commit package-lock.json, "
     "provide Docker fallback, test install on clean machine before release."),
    (8, "Insufficient automated test coverage leading to regressions in API and UI",
     "TE", 30, 2,
     "Maintain Node test runner suite for API and Playwright suite for UI, run on "
     "every PR, target >70% coverage on routes, add tests before bug-fix merges."),
]

# Sorted descending by probability
SORTED_RISKS = sorted(RISKS, key=lambda r: -r[3])

# Risk Information Sheets for top 3 risks (by probability)
RIS_DETAILS = {
    1: {  # Schedule slippage
        "subs": [
            "Initial effort estimation may be optimistic due to unfamiliar tech stack.",
            "Inter-module dependencies (auth -> activities -> twin status) cause cascading delays.",
            "Team members balancing other coursework reduces effective hours.",
        ],
        "mitigation": [
            "Add 15-20% buffer time per sprint.",
            "Break features into measurable user stories.",
            "Track milestones in a shared Gantt chart and review weekly.",
            "Use GitHub Projects/Issues for visibility.",
        ],
        "contingency": [
            "Re-prioritize backlog and defer non-critical features.",
            "Reallocate tasks to less-loaded members.",
            "Inform mentor and request scope adjustment if needed.",
            "Trim optional UI polish to protect core deliverables.",
        ],
        "trigger": "Two consecutive sprint goals missed, or critical path task delayed by > 3 days.",
        "status": "Monitoring in progress; weekly reviews active.",
        "assigned": "Project Manager / Scrum Master",
        "originator": "Risk Management Team",
    },
    2: {  # Twin computation accuracy
        "subs": [
            "Health-score formula not validated against domain data.",
            "Edge cases (missing days, zero activity) produce skewed twin status.",
            "Aggregation queries on SQLite may round or truncate incorrectly.",
        ],
        "mitigation": [
            "Document twin-status formula and review with mentor.",
            "Add unit tests in tests/api.test.js for aggregation endpoints.",
            "Seed the DB with edge-case data via src/seed.js and assert outputs.",
            "Peer-review src/routes/twin.js logic before merge.",
        ],
        "contingency": [
            "Roll back to previous formula version via git.",
            "Hot-patch the calculation and re-run regression tests.",
            "Notify users that twin status is being recalibrated.",
        ],
        "trigger": "Twin status output deviates > 10% from expected reference values, or user-reported anomalies.",
        "status": "Mitigation in design phase.",
        "assigned": "Backend Developer",
        "originator": "Risk Management Team",
    },
    3: {  # Team experience
        "subs": [
            "Members new to Express middleware and async patterns.",
            "Playwright/Node test runner setup unfamiliar to most members.",
            "SQL query optimization knowledge limited.",
        ],
        "mitigation": [
            "Run weekly knowledge-sharing sessions.",
            "Pair-program on complex modules (auth, twin logic).",
            "Maintain a project README with setup and conventions.",
            "Adopt ESLint + code-review checklist.",
        ],
        "contingency": [
            "Consult mentor or senior peers when blocked > 1 day.",
            "Use well-documented libraries instead of custom code where possible.",
            "Reassign tasks to members with relevant strength.",
        ],
        "trigger": "PRs repeatedly rejected for the same class of issue, or task blocked > 24 hours.",
        "status": "Mitigation actions initiated.",
        "assigned": "Tech Lead",
        "originator": "Risk Management Team",
    },
}


# ---------- helpers ----------

def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return p


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')
        for r in p.runs:
            r.font.size = Pt(11)


def make_risk_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=6)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    headers = ["Risk ID", "Risks", "Category", "Probability", "Impact", "RMMM"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        set_cell_shading(hdr[i], "1F3A5F")
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        rid, risk, cat, prob, impact, rmmm = row
        cells = table.rows[ri].cells
        cells[0].text = str(rid)
        cells[1].text = risk
        cells[2].text = cat
        cells[3].text = f"{prob}%"
        cells[4].text = str(impact)
        cells[5].text = rmmm
        for c in cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    # column widths
    widths = [Cm(1.6), Cm(5.5), Cm(1.8), Cm(2.2), Cm(1.6), Cm(6.0)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
    return table


def add_ris(doc, risk_tuple, details):
    rid, risk, cat, prob, impact, rmmm = risk_tuple
    add_heading(doc, f"Risk Information Sheet (RIS) – Risk {rid} ({prob}%)", level=2)

    # top info table
    info = doc.add_table(rows=2, cols=4)
    info.style = 'Light Grid Accent 1'
    info.rows[0].cells[0].text = "Risk ID"
    info.rows[0].cells[1].text = str(rid)
    info.rows[0].cells[2].text = "Date"
    info.rows[0].cells[3].text = DATE
    info.rows[1].cells[0].text = "Probability"
    info.rows[1].cells[1].text = f"{prob}%"
    info.rows[1].cells[2].text = "Impact"
    info.rows[1].cells[3].text = str(impact)
    for row in info.rows:
        for i, c in enumerate(row.cells):
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if i % 2 == 0:
                        r.bold = True

    add_para(doc, "Description", bold=True)
    add_para(doc, risk)

    add_para(doc, "Refinement & Context", bold=True)
    add_bullets(doc, [f"Sub-condition {i+1}: {s}" for i, s in enumerate(details["subs"])])

    add_para(doc, "Mitigation & Monitoring Strategies", bold=True)
    add_bullets(doc, details["mitigation"])

    add_para(doc, "Contingency Plan and Management", bold=True)
    add_bullets(doc, details["contingency"])

    add_para(doc, "Trigger", bold=True)
    add_para(doc, details["trigger"])

    add_para(doc, "Status", bold=True)
    add_para(doc, details["status"])

    meta = doc.add_table(rows=1, cols=2)
    meta.style = 'Light Grid Accent 1'
    meta.rows[0].cells[0].text = f"Assigned To: {details['assigned']}"
    meta.rows[0].cells[1].text = f"Originator: {details['originator']}"
    for c in meta.rows[0].cells:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                r.bold = True

    doc.add_paragraph()


# ---------- build doc ----------

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run(f"RMMM Plan – {PROJECT_NAME}")
trun.bold = True
trun.font.size = Pt(18)
trun.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
srun = sub.add_run("Risk Mitigation, Monitoring, and Management Plan")
srun.italic = True
srun.font.size = Pt(12)

doc.add_paragraph()

# Project Overview
add_heading(doc, "1. Project Overview", level=1)
add_para(doc,
    f"{PROJECT_NAME} is a web-based health activity tracking system built on the "
    "concept of a Digital Twin. It is implemented using Node.js with the Express "
    "framework, persisted in a SQLite database (better-sqlite3), and served with a "
    "lightweight HTML/CSS/JS frontend. Users can register/log in, log daily "
    "activities (steps, sleep, calories, workouts), and view a continuously updated "
    "digital-twin status that reflects their current health profile. The system is "
    "tested via the Node.js test runner for API endpoints and Playwright for UI flows."
)

add_para(doc, "Key components:", bold=True)
add_bullets(doc, [
    "REST API routes: /api/users, /api/users/:id/activities, /api/users/:id/twin-status",
    "SQLite store (health.db) with users and activities models",
    "Static frontend served from /public (index.html, login.html, JS modules)",
    "Automated tests under /tests (API) and /tests-ui (Playwright)",
    "Environment configuration via .env",
])

# Step 1
add_heading(doc, "2. Step 1 – Risk Table", level=1)
add_para(doc,
    "The following table lists the identified risks for the project, along with "
    "their category, estimated probability, impact, and the corresponding RMMM action."
)
make_risk_table(doc, RISKS)

# Step 2
add_heading(doc, "3. Step 2 – Risks Sorted by Probability (Descending)", level=1)
make_risk_table(doc, SORTED_RISKS)

# Step 3
add_heading(doc, "4. Step 3 – Risk Information Sheets (RIS)", level=1)
add_para(doc,
    "Detailed Risk Information Sheets are presented below for the top three risks "
    "(highest probability of occurrence). The same template applies to the remaining "
    "risks in the risk register."
)

# Map id -> risk tuple
risk_by_id = {r[0]: r for r in RISKS}
for rid in [1, 2, 3]:
    add_ris(doc, risk_by_id[rid], RIS_DETAILS[rid])

# Conclusion
add_heading(doc, "5. Conclusion", level=1)
add_para(doc,
    "This RMMM Plan identifies the most significant risks for the "
    f"{PROJECT_NAME} and defines clear mitigation, monitoring, and contingency "
    "actions for each. The risk register will be reviewed at the end of every "
    "sprint, and Risk Information Sheets will be updated whenever a risk's "
    "probability, impact, or status changes materially."
)

out_path = "RMMM_Plan_Digital_Twin_Health_Tracker.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
